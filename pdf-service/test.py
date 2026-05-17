"""
Local test script for PDF service.
Starts the server, sends test files, verifies results.

Usage:
    python test.py              # test PDF→DOCX
    python test.py --docx2pdf   # also test DOCX→PDF (needs LibreOffice)
"""

import os, sys, time, json, urllib.request, urllib.error, base64, tempfile, shutil
from pathlib import Path
import subprocess
import threading

HOST = "http://127.0.0.1:8000"
SAMPLE_DIR = Path(tempfile.gettempdir()) / "pdf_service_test"


def log(msg: str):
    print(f"[test] {msg}")


def ensure_sample_files():
    SAMPLE_DIR.mkdir(parents=True, exist_ok=True)

    pdf_path = SAMPLE_DIR / "sample.pdf"
    if not pdf_path.exists():
        log("Creating sample PDF with PyMuPDF...")
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 100), "Hello PDF→DOCX!", fontsize=20)
        page.insert_text((50, 150), "中文测试文本", fontsize=16)
        doc.save(str(pdf_path))
        doc.close()
        log(f"  Created {pdf_path} ({pdf_path.stat().st_size} bytes)")

    docx_path = SAMPLE_DIR / "sample.docx"
    if not docx_path.exists():
        log("Creating sample DOCX with python-docx...")
        from docx import Document
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("Hello DOCX→PDF!")
        doc.add_paragraph("中文段落测试")
        doc.add_paragraph("This is a paragraph with some regular text.")
        doc.save(str(docx_path))
        log(f"  Created {docx_path} ({docx_path.stat().st_size} bytes)")

    return pdf_path, docx_path


def wait_for_health(url: str, timeout: int = 30):
    deadline = time.time() + timeout
    while time.time() < deadline:
        try:
            resp = urllib.request.urlopen(f"{url}/health", timeout=5)
            if resp.status == 200:
                log("Server is healthy")
                return True
        except Exception:
            pass
        time.sleep(1)
    return False


def upload_and_convert(file_path: Path, from_fmt: str, to_fmt: str) -> bytes:
    """Upload file to /convert, poll /status, download result. Return raw bytes."""
    with open(file_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("ascii")

    body = json.dumps({
        "file_base64": b64,
        "filename": file_path.name,
        "from_fmt": from_fmt,
        "to_fmt": to_fmt,
    }).encode("utf-8")

    req = urllib.request.Request(
        f"{HOST}/convert",
        data=body,
        headers={"Content-Type": "application/json"},
    )
    resp = urllib.request.urlopen(req, timeout=60)
    result = json.loads(resp.read())
    job_id = result["job_id"]
    log(f"Job submitted: {job_id}")

    # Poll until done
    for _ in range(60):
        time.sleep(2)
        status_resp = urllib.request.urlopen(f"{HOST}/status/{job_id}", timeout=30)
        status = json.loads(status_resp.read())
        if status["status"] == "done":
            log(f"Job done, downloading...")
            dl = urllib.request.urlopen(f"{HOST}/download/{status['result']}", timeout=60)
            return dl.read()
        elif status["status"] == "error":
            raise RuntimeError(f"Conversion failed: {status.get('error', 'unknown')}")
        else:
            log(f"  Status: {status['status']}")

    raise TimeoutError("Conversion did not complete within timeout")


def verify_pdf(data: bytes):
    assert data[:4] == b"%PDF", f"Bad PDF header: {data[:20]}"
    log(f"  Valid PDF, size={len(data)} bytes")


def verify_docx(data: bytes):
    assert data[:2] == b"PK", f"Bad DOCX (ZIP) header: {data[:20]}"
    log(f"  Valid DOCX, size={len(data)} bytes")


def start_server():
    """Start uvicorn as subprocess, return process handle."""
    env = os.environ.copy()
    env["PORT"] = "8000"
    proc = subprocess.Popen(
        [sys.executable, "-m", "uvicorn", "main:app", "--host", "127.0.0.1", "--port", "8000"],
        cwd=str(Path(__file__).parent),
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        env=env,
    )
    return proc


def main():
    test_docx2pdf = "--docx2pdf" in sys.argv

    pdf_path, docx_path = ensure_sample_files()

    log("Starting server...")
    server = start_server()
    try:
        if not wait_for_health(HOST):
            out, err = server.communicate(timeout=5)
            log(f"Server failed to start. stdout={out.decode()}, stderr={err.decode()}")
            sys.exit(1)

        # ── Test 1: PDF → DOCX ──
        log("=" * 50)
        log("Test 1: PDF → DOCX")
        log("=" * 50)
        try:
            result = upload_and_convert(pdf_path, "pdf", "docx")
            verify_docx(result)
            log("PASS: PDF→DOCX")
        except Exception as e:
            log(f"FAIL: PDF→DOCX — {e}")
            raise

        # ── Test 2: DOCX → PDF (if requested or LO available) ──
        if test_docx2pdf:
            log("=" * 50)
            log("Test 2: DOCX → PDF")
            log("=" * 50)
            try:
                result = upload_and_convert(docx_path, "docx", "pdf")
                verify_pdf(result)
                log("PASS: DOCX→PDF")
            except Exception as e:
                log(f"FAIL: DOCX→PDF — {e}")
                raise
        else:
            log("Skipping DOCX→PDF test (use --docx2pdf to enable)")

        log("=" * 50)
        log("All tests passed!")

    finally:
        server.terminate()
        server.wait(timeout=5)


if __name__ == "__main__":
    main()
