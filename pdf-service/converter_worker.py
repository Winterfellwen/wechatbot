"""
Isolated subprocess worker for PDF conversion.
- PDF→DOCX: splits large PDFs into page-range chunks, converts each with pdf2docx,
  then merges DOCX outputs. Memory per chunk stays low.
- DOCX→PDF: via LibreOffice subprocess.
Usage: python converter_worker.py <input_path> <output_path> <from_fmt> <to_fmt>
"""

import os
import sys
import shutil
import gc
import subprocess
import uuid
import time
import struct
import zlib
import zipfile
import math
from pathlib import Path

UPLOAD_DIR = Path(os.environ.get("PDF_TEMP_DIR", "/tmp")) / "pdf-service"
CHUNK_MAX_PAGES = int(os.environ.get("CHUNK_MAX_PAGES", "15"))  # pages per chunk
CHUNK_MAX_SIZE = int(os.environ.get("CHUNK_MAX_SIZE", "15"))  # MB per chunk estimate
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def safe_unlink(path: Path):
    try:
        if path.exists():
            path.unlink(missing_ok=True)
    except Exception:
        pass


def find_libreoffice() -> str | None:
    env_path = os.environ.get("LIBREOFFICE_PATH")
    if env_path and os.path.isfile(env_path):
        return env_path
    candidates = [
        "/usr/bin/libreoffice", "/usr/bin/soffice",
        "/usr/local/bin/libreoffice", "/usr/local/bin/soffice",
        "/snap/bin/libreoffice",
    ]
    for p in candidates:
        if os.path.isfile(p):
            return p
    for bin_name in ("libreoffice", "soffice"):
        try:
            r = subprocess.run(["which", bin_name], capture_output=True, text=True, timeout=5)
            if r.returncode == 0 and r.stdout.strip():
                return r.stdout.strip()
        except Exception:
            pass
    return None


# ============================================================
# DOCX merge
# ============================================================
def merge_docx(chunk_paths: list[Path], output_path: Path):
    """Merge multiple DOCX files into one by copying content."""
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    merged = Document()
    # Remove default empty paragraph
    if merged.paragraphs:
        p = merged.paragraphs[0]._element
        p.getparent().remove(p)

    for i, chunk_path in enumerate(chunk_paths):
        print(f"[worker]  Merging chunk {i+1}/{len(chunk_paths)}: {chunk_path.name}", flush=True)
        chunk = Document(str(chunk_path))
        for element in chunk.element.body:
            merged.element.body.append(element)
        chunk = None
        gc.collect()

    merged.save(str(output_path))
    merged = None
    gc.collect()
    print(f"[worker]  Merged DOCX: {output_path.name} size={output_path.stat().st_size}", flush=True)


# ============================================================
# PDF → DOCX (chunked, low memory per chunk)
# ============================================================
def pdf_to_docx(input_path: Path, output_path: Path):
    try:
        import fitz
    except ImportError:
        raise RuntimeError("PyMuPDF (fitz) required")
    try:
        from pdf2docx import Converter
    except ImportError:
        raise RuntimeError("pdf2docx required")

    pdf = fitz.open(str(input_path))
    num_pages = len(pdf)
    file_mb = input_path.stat().st_size / (1024 * 1024)
    print(f"[worker] PDF: {num_pages} pages, {file_mb:.1f}MB", flush=True)

    # Determine chunk size: max 15 pages OR ~15MB per chunk
    estimated_mb_per_page = max(file_mb / max(num_pages, 1), 0.1)
    pages_per_chunk = max(1, min(CHUNK_MAX_PAGES, int(CHUNK_MAX_SIZE / estimated_mb_per_page)))
    num_chunks = math.ceil(num_pages / pages_per_chunk)

    print(f"[worker] Splitting into {num_chunks} chunk(s) of ~{pages_per_chunk} pages each", flush=True)
    pdf.close()
    pdf = None
    gc.collect()

    if num_chunks == 1:
        # Small PDF — convert directly
        print(f"[worker] Single chunk, converting directly...", flush=True)
        _convert_pdf_chunk(input_path, output_path, 0, num_pages)
        print(f"[worker] Done: {output_path.name} size={output_path.stat().st_size}", flush=True)
        return

    # Large PDF — split, convert chunks, merge
    chunk_files = []
    try:
        for chunk_idx in range(num_chunks):
            start_page = chunk_idx * pages_per_chunk
            end_page = min(start_page + pages_per_chunk, num_pages)
            chunk_pdf = UPLOAD_DIR / f"{input_path.stem}_chunk{chunk_idx}.pdf"
            chunk_docx = UPLOAD_DIR / f"{input_path.stem}_chunk{chunk_idx}.docx"

            # Extract page range to temp PDF using fitz (zero-copy)
            src = fitz.open(str(input_path))
            dst = fitz.open()
            dst.insert_pdf(src, from_page=start_page, to_page=end_page - 1)
            dst.save(str(chunk_pdf), garbage=4, deflate=True)
            dst.close()
            src.close()
            gc.collect()
            print(f"[worker] Chunk {chunk_idx+1}: pages {start_page+1}-{end_page} → {chunk_pdf.name} ({chunk_pdf.stat().st_size//1024}KB)", flush=True)

            _convert_pdf_chunk(chunk_pdf, chunk_docx, start_page, end_page)
            chunk_files.append(chunk_docx)

        # Merge all chunks
        print(f"[worker] Merging {len(chunk_files)} DOCX chunks...", flush=True)
        merge_docx(chunk_files, output_path)

    finally:
        for f in chunk_files:
            safe_unlink(f)
        for f in UPLOAD_DIR.glob(f"{input_path.stem}_chunk*.pdf"):
            safe_unlink(f)


def _convert_pdf_chunk(in_path: Path, out_path: Path, start_page: int, end_page: int):
    """Convert a single PDF chunk with pdf2docx."""
    from pdf2docx import Converter
    print(f"[worker]  Converting chunk pages {start_page+1}-{end_page}...", flush=True)
    cv = Converter(str(in_path))
    cv.convert(str(out_path))
    cv.close()
    gc.collect()
    size_mb = out_path.stat().st_size / (1024 * 1024) if out_path.exists() else 0
    print(f"[worker]  Chunk done: {size_mb:.1f}MB", flush=True)


# ============================================================
# DOCX → PDF (via LibreOffice subprocess)
# ============================================================
def kill_libreoffice():
    try:
        subprocess.run(["pkill", "-f", "libreoffice"], capture_output=True, timeout=10)
        subprocess.run(["pkill", "-f", "soffice.bin"], capture_output=True, timeout=10)
    except Exception:
        pass


def docx_to_pdf(input_path: Path, output_path: Path):
    libreoffice_bin = find_libreoffice()
    if not libreoffice_bin:
        raise RuntimeError("LibreOffice not found")

    job_tag = uuid.uuid4().hex[:8]
    lo_home = UPLOAD_DIR / f"lo_home_{job_tag}"
    lo_home.mkdir(parents=True, exist_ok=True)
    tmp_out = UPLOAD_DIR / f"lo_out_{job_tag}"
    tmp_out.mkdir(exist_ok=True)

    try:
        lo_env = os.environ.copy()
        lo_env.pop("SAL_USE_VCLPLUGIN", None)
        lo_env["HOME"] = str(lo_home)
        lo_env["SAL_DISABLE_OPENGL_CHECK"] = "1"
        lo_env["SAL_VIDEO_DISABLE_ACCELERATE"] = "1"
        lo_env["LIBREOFFICE_MEMORY_MULTIPLIER"] = "0.3"
        lo_env["OOO_DISABLE_RECOVERY"] = "1"

        cmd = [
            libreoffice_bin,
            f"-env:UserInstallation=file://{lo_home}",
            "--headless", "--norestore", "--nofirststartwizard",
            "--convert-to", "pdf:writer_pdf_Export",
            "--outdir", str(tmp_out),
            str(input_path),
        ]

        kill_libreoffice()
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=300, env=lo_env)
        if result.returncode != 0:
            kill_libreoffice()
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300, env=lo_env)

        pdf_files = list(tmp_out.glob("*.pdf"))
        if not pdf_files:
            raise RuntimeError(f"LibreOffice produced no PDF. stderr: {(result.stderr or '')[:1000]}")
        lo_pdf_path = pdf_files[0]
        if lo_pdf_path.stat().st_size == 0:
            raise RuntimeError("LibreOffice produced an empty PDF")

        shutil.copy2(lo_pdf_path, output_path)
    finally:
        shutil.rmtree(tmp_out, ignore_errors=True)
        shutil.rmtree(lo_home, ignore_errors=True)
        kill_libreoffice()
        gc.collect()


# ============================================================
def main():
    if len(sys.argv) != 5:
        print("Usage: converter_worker.py <input_path> <output_path> <from_fmt> <to_fmt>", file=sys.stderr)
        sys.exit(1)

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    from_fmt = sys.argv[3]
    to_fmt = sys.argv[4]

    if not input_path.exists():
        print(f"Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    try:
        start = time.time()
        print(f"[worker] start {from_fmt}→{to_fmt} input={input_path.name} size={input_path.stat().st_size//1024}KB", flush=True)

        if from_fmt == "pdf" and to_fmt == "docx":
            pdf_to_docx(input_path, output_path)
        elif from_fmt == "docx" and to_fmt == "pdf":
            docx_to_pdf(input_path, output_path)
        else:
            print(f"Unsupported conversion: {from_fmt} -> {to_fmt}", file=sys.stderr)
            sys.exit(1)

        elapsed = time.time() - start
        print(f"[worker] done in {elapsed:.1f}s output={output_path.name} size={output_path.stat().st_size}", flush=True)
        sys.exit(0)
    except Exception as e:
        print(f"[worker] error: {e}", file=sys.stderr, flush=True)
        sys.exit(1)


if __name__ == "__main__":
    main()
