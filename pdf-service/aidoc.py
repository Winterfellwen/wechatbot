import os
import base64
import uuid
import json
import requests
import re
from pathlib import Path
from fastapi import APIRouter, HTTPException, Form
from fastapi.responses import JSONResponse
from pydantic import BaseModel
import fitz

router = APIRouter()

UPLOAD_DIR = Path("/tmp/aidoc")
UPLOAD_DIR.mkdir(exist_ok=True)

OUTPUT_DIR = Path("/tmp/aidoc-output")
OUTPUT_DIR.mkdir(exist_ok=True)

NODE_API_URL = os.environ.get("NODE_API_URL", "http://localhost:3000")


class ConvertToHtmlRequest(BaseModel):
    file_base64: str
    filename: str


class AIReviewRequest(BaseModel):
    html_content: str
    instructions: str = "检查HTML内容，修复布局问题，调整图片大小，确保格式正确。"


class ExportRequest(BaseModel):
    html_content: str = ""
    format: str = "pdf"
    job_id: str = ""


def convert_pdf_to_html(file_data: bytes, filename: str) -> str:
    """将PDF转换为简单的HTML用于预览（文字+图片，简化版）"""
    import fitz
    import base64

    pages_html = []

    try:
        doc = fitz.open(stream=file_data, filetype="pdf")
    except Exception as e:
        raise Exception(f"Cannot open PDF: {str(e)}")

    for page_num, page in enumerate(doc):
        w = int(page.rect.width)
        h = int(page.rect.height)

        page_content = [f'<div class="page" style="width:{w}px;height:{h}px;position:relative;background:white;margin:10px auto;box-shadow:0 2px 8px rgba(0,0,0,0.15);">']

        # 提取图片
        try:
            for img in page.get_images():
                try:
                    xref = img[0]
                    img_info = page.parent.extract_image(xref)
                    if img_info and "image" in img_info:
                        b64 = base64.b64encode(img_info["image"]).decode()
                        ext = img_info.get("ext", "png")
                        bbox = img[1]
                        x, y = int(bbox[0]), int(bbox[1])
                        ww, hh = int(bbox[2] - bbox[0]), int(bbox[3] - bbox[1])
                        page_content.append(f'<img src="data:image/{ext};base64,{b64}" style="position:absolute;left:{x}px;top:{y}px;width:{ww}px;height:{hh}px;" />')
                except:
                    pass
        except:
            pass

        # 提取文字（按行输出，不使用绝对定位，简单堆叠）
        try:
            text = page.get_text("text")
            if text.strip():
                text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                # 使用pre标签简单显示
                page_content.append(f'<pre style="margin:0;padding:10px;font-size:12px;white-space:pre-wrap;word-wrap:break-word;">{text}</pre>')
        except:
            pass

        page_content.append('</div>')
        pages_html.append('\n'.join(page_content))

    doc.close()

    html = f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width,initial-scale=1.0,user-scalable=yes">
    <style>
        body {{ font-family: sans-serif; background: #f0f0f0; margin: 0; padding: 10px; }}
        .page {{ position: relative; overflow: hidden; }}
    </style>
</head>
<body>
{chr(10).join(pages_html)}
</body>
</html>'''

    return html


def convert_docx_to_html(file_data: bytes) -> str:
    """使用mammoth将docx转换为HTML"""
    import mammoth

    result = mammoth.convert_to_html({"file_data": file_data})
    html_content = result.get("value", "")

    # 构建完整的HTML文档，添加缩放功能
    html = []
    html.append('<!DOCTYPE html><html><head><meta charset="utf-8">')
    html.append('<meta name="viewport" content="width=device-width,initial-scale=1.0">')
    html.append('<style>')
    html.append('* { box-sizing: border-box; }')
    html.append('body { font-family: sans-serif; margin: 0; padding: 10px; background: #f0f0f0; }')
    html.append('.container { transform-origin: top left; transition: transform 0.3s; }')
    html.append('.page { background: white; margin: 10px auto; padding: 20px; box-shadow: 0 2px 8px rgba(0,0,0,0.15); }')
    html.append('img { max-width: 100%; height: auto; }')
    html.append('</style>')
    html.append('<script>')
    html.append('function scaleToFit() {')
    html.append('  var c = document.querySelector(".container"); if (!c) return;')
    html.append('  var p = c.querySelector(".page"); if (!p) return;')
    html.append('  var s = (window.innerWidth - 20) / p.offsetWidth;')
    html.append('  if (s > 1) s = 1;')
    html.append('  c.style.transform = "scale(" + s + ")"; c.style.width = (100/s) + "%";')
    html.append('}')
    html.append('window.onload = scaleToFit; window.onresize = scaleToFit;')
    html.append('</script></head><body>')
    html.append('<div class="container"><div class="page">')
    html.append(html_content)
    html.append('</div></div></body></html>')

    return '\n'.join(html)


@router.post("/convert-to-html")
async def convert_to_html(req: ConvertToHtmlRequest):
    """将PDF或DOCX文件转换为HTML5"""
    try:
        file_data = base64.b64decode(req.file_base64)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid base64: {str(e)}")

    filename = req.filename.lower()
    job_id = uuid.uuid4().hex

    OUTPUT_DIR.mkdir(exist_ok=True)

    try:
        if filename.endswith('.pdf'):
            html = convert_pdf_to_html(file_data, filename)
        elif filename.endswith('.docx'):
            html = convert_docx_to_html(file_data)
        elif filename.endswith('.doc'):
            html = convert_docx_to_html(file_data)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Use PDF, DOCX, or DOC")

        output_file = OUTPUT_DIR / f"{job_id}.html"
        output_file.write_text(html, encoding='utf-8')

        return {
            "job_id": job_id,
            "status": "done",
            "html": html,
            "html_url": f"/aidoc/html/{job_id}.html"
        }

    except Exception as e:
        import traceback
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}\n{traceback.format_exc()}")


@router.post("/ai-review")
async def ai_review(req: AIReviewRequest):
    """使用AI修正HTML内容"""
    try:
        prompt = f"""你是一个HTML文档修正专家。请检查以下HTML内容，找出并修复以下问题：
1. 布局问题（如元素重叠、溢出）
2. 图片大小问题（过大、过小、变形）
3. 格式问题（字体、颜色、间距）
4. 语法错误（如未闭合的标签）

用户指令: {req.instructions}

请直接返回修正后的HTML代码（包含完整的<!DOCTYPE html>和<html>标签），不要添加任何解释或markdown格式。"""

        messages = [
            {"role": "system", "content": "你是一个专业的HTML文档修正助手。"},
            {"role": "user", "content": f"{prompt}\n\n原始HTML:\n{req.html_content}"}
        ]

        node_response = requests.post(
            f"{NODE_API_URL}/api/chat",
            json={"messages": messages},
            timeout=120
        )

        if node_response.status_code != 200:
            raise HTTPException(status_code=500, detail="AI service unavailable")

        result = node_response.json()

        if 'choices' in result and len(result['choices']) > 0:
            corrected_html = result['choices'][0]['message']['content']

            corrected_html = re.sub(r'^```html\n?', '', corrected_html)
            corrected_html = re.sub(r'^```\n?$', '', corrected_html)
            corrected_html = corrected_html.strip()

            return {
                "status": "done",
                "original_html": req.html_content,
                "corrected_html": corrected_html
            }
        else:
            raise HTTPException(status_code=500, detail="AI response invalid")

    except requests.RequestException as e:
        raise HTTPException(status_code=500, detail=f"AI service error: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Review failed: {str(e)}")


@router.post("/export")
async def export_html(req: ExportRequest):
    """将HTML导出为PDF、DOCX或DOC格式"""
    import mammoth
    import weasyprint
    from docx import Document
    from io import BytesIO

    OUTPUT_DIR.mkdir(exist_ok=True)

    html = req.html_content
    job_id = req.job_id

    # 如果没有html_content但有job_id，从文件读取
    if not html and job_id:
        html_file = OUTPUT_DIR / f"{job_id}.html"
        if html_file.exists():
            html = html_file.read_text(encoding='utf-8')

    if not html:
        raise HTTPException(status_code=400, detail="No HTML content provided")

    fmt = req.format.lower()
    if not job_id:
        job_id = uuid.uuid4().hex

    if fmt not in ['pdf', 'docx', 'doc']:
        raise HTTPException(status_code=400, detail="Format must be pdf, docx, or doc")

    try:
        if fmt == 'pdf':
            output_path = OUTPUT_DIR / f"{job_id}.pdf"
            weasyprint.HTML(string=html).write_pdf(str(output_path))

            with open(output_path, 'rb') as f:
                file_data = base64.b64encode(f.read()).decode()

            return {
                "job_id": job_id,
                "format": "pdf",
                "file_base64": file_data,
                "filename": f"exported_{job_id}.pdf"
            }

        elif fmt in ['docx', 'doc']:
            result = mammoth.convert_to_html({"file_data": html.encode('utf-8')})
            html_content = result.get("value", "")

            doc = Document()
            doc.add_heading('Document', 0)

            temp_html = f"<html><body>{html_content}</body></html>"
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(temp_html, 'html.parser')

            for tag in soup.find_all(['p', 'h1', 'h2', 'h3', 'li']):
                if tag.name == 'h1':
                    doc.add_heading(tag.get_text(), level=1)
                elif tag.name == 'h2':
                    doc.add_heading(tag.get_text(), level=2)
                elif tag.name == 'h3':
                    doc.add_heading(tag.get_text(), level=3)
                elif tag.name == 'li':
                    doc.add_paragraph(tag.get_text(), style='List Bullet')
                else:
                    doc.add_paragraph(tag.get_text())

            if fmt == 'doc':
                docx_path = OUTPUT_DIR / f"{job_id}.docx"
                doc.save(str(docx_path))

                docx_file = docx_path.read_bytes()

                import zipfile
                from pathlib import Path as PathLib
                doc_path = OUTPUT_DIR / f"{job_id}.doc"
                with zipfile.ZipFile(BytesIO(docx_file)) as z_in:
                    with zipfile.ZipFile(doc_path, 'w', zipfile.ZIP_DEFLATED) as z_out:
                        for item in z_in.infolist():
                            if item.filename != 'word/document.xml':
                                z_out.writestr(item, z_in.read(item.filename))
                            else:
                                content = z_in.read(item.filename).decode('utf-8')
                                content = content.replace('w: vals="1"', 'w: vals="0 1 2"')
                                content = content.replace('w:fldCharType="end"', 'w:fldCharType="begin"')
                                z_out.writestr(item, content.encode('utf-8'))

                with open(doc_path, 'rb') as f:
                    file_data = base64.b64encode(f.read()).decode()
                return {
                    "job_id": job_id,
                    "format": "doc",
                    "file_base64": file_data,
                    "filename": f"exported_{job_id}.doc"
                }
            else:
                docx_path = OUTPUT_DIR / f"{job_id}.docx"
                doc.save(str(docx_path))

                with open(docx_path, 'rb') as f:
                    file_data = base64.b64encode(f.read()).decode()

                return {
                    "job_id": job_id,
                    "format": "docx",
                    "file_base64": file_data,
                    "filename": f"exported_{job_id}.docx"
                }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


@router.get("/html/{filename}")
async def get_html(filename: str):
    """获取HTML文件"""
    from fastapi.responses import Response

    OUTPUT_DIR.mkdir(exist_ok=True)
    path = OUTPUT_DIR / filename
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"File not found: {filename}")

    try:
        content = path.read_text(encoding='utf-8')
        return Response(content=content, media_type="text/html")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Read error: {str(e)}")


@router.get("/edit/{filename}")
async def get_editable_html(filename: str):
    """获取可编辑的HTML页面"""
    from fastapi.responses import Response

    OUTPUT_DIR.mkdir(exist_ok=True)
    path = OUTPUT_DIR / filename
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"File not found: {filename}")

    try:
        content = path.read_text(encoding='utf-8')

        # 注入编辑功能
        editable_html = content.replace('<head>', '''<head>
    <style>
        [contenteditable="true"] { outline: 2px dashed #1890ff; padding: 4px; }
        [contenteditable="true"]:hover { background: #f0f7ff; }
        .save-btn { position: fixed; bottom: 20px; right: 20px; padding: 12px 24px; 
            background: #1890ff; color: white; border: none; border-radius: 8px; 
            font-size: 16px; z-index: 9999; box-shadow: 0 4px 12px rgba(0,0,0,0.15); }
    </style>
    <script>
        function saveContent() {
            document.body.style.cursor = "wait";
            var html = document.body.innerHTML;
            // 移除script标签防止XSS
            html = html.replace(/<script[\s\S]*?<\/script>/gi, "");
            // 调用导出API保存
            fetch("/aidoc/export", {
                method: "POST",
                headers: {"Content-Type": "application/json"},
                body: JSON.stringify({html_content: html, format: "pdf", job_id: "%s"})
            }).then(r => r.json()).then(data => {
                alert("保存成功！");
                document.body.style.cursor = "default";
            }).catch(e => {
                alert("保存失败: " + e);
                document.body.style.cursor = "default";
            });
        }
    </script>''' % filename.replace('.html', ''))

        # 添加保存按钮
        editable_html = editable_html.replace('</body>', 
            '<button class="save-btn" onclick="saveContent()">💾 保存修改</button></body>')

        return Response(content=editable_html, media_type="text/html")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Read error: {str(e)}")